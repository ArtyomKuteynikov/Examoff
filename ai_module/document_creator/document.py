from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import re

from docx.styles.style import ParagraphStyle


def add_paragraph_format(style: ParagraphStyle) -> None:
    """
    Настраивает абзац у стиля.
        Междустрочный интервал – полуторный.
        Отступ первой строки основного текста 1,25 см.
        Выравнивание основного текста – по ширине.

    :param style: Стиль для настройки.
    """
    style.paragraph_format.space_before = Mm(0)
    style.paragraph_format.space_after = Mm(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.first_line_indent = Cm(1.25)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def init_main_text_style(document: Document) -> None:
    """
    Стиль для `Основной текст`.
        Шрифт `Основной текст` – Times New Roman, 14.

    :param document: Word документ.
    """
    main_text_style = document.styles.add_style('Main Text Style', WD_STYLE_TYPE.PARAGRAPH)

    # Шрифт
    main_text_style.font.name = 'Times New Roman'
    main_text_style.font.size = Pt(14)

    # Абзац
    add_paragraph_format(main_text_style)


def init_chapter_headings_style(document: Document) -> None:
    """
    Стиль для `Заголовки глав`.
        Шрифт `Заголовки глав` – Arial, 16, полужирный, курсив.

    :param document: Word документ.
    """
    main_text_style = document.styles.add_style('Chapter Headings Style', WD_STYLE_TYPE.PARAGRAPH)

    # Шрифт
    main_text_style.font.name = 'Arial'
    main_text_style.font.size = Pt(16)
    main_text_style.font.bold = True
    main_text_style.font.italic = True

    # Абзац
    add_paragraph_format(main_text_style)


def init_paragraph_headings_style(document: Document) -> None:
    """
    Стиль для `Заголовки параграфов`.
        Шрифт `Заголовки параграфов` – Times New Roman, 16, полужирный.

    :param document: Word документ.
    """
    main_text_style = document.styles.add_style('Paragraph Headings Style', WD_STYLE_TYPE.PARAGRAPH)

    # Шрифт
    main_text_style.font.name = 'Times New Roman'
    main_text_style.font.size = Pt(16)
    main_text_style.font.bold = True

    # Абзац
    add_paragraph_format(main_text_style)


def init_item_headings_style(document: Document) -> None:
    """
    Стиль для `Заголовки пунктов`.
        Шрифт `Заголовки пунктов` – Times New Roman, 14, полужирный.

    :param document: Word документ.
    """
    main_text_style = document.styles.add_style('Item Headings Style', WD_STYLE_TYPE.PARAGRAPH)

    # Шрифт
    main_text_style.font.name = 'Times New Roman'
    main_text_style.font.size = Pt(14)
    main_text_style.font.bold = True

    # Абзац
    add_paragraph_format(main_text_style)


def init_styles(document: Document) -> None:
    """
    Добавляет стили форматирования для Document.

    :param document: Word документ.
    """
    # Параметры страницы: правое поле – 3 см, левое поле – 1,5 см, верхнее и нижнее поля – 2 см.
    # http://genhis.philol.msu.ru/trebovaniya-k-diplomnoj-rabote/
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(3)

    init_main_text_style(document)
    init_chapter_headings_style(document)
    init_paragraph_headings_style(document)
    init_item_headings_style(document)


if __name__ == '__main__':
    my_document = Document()
    init_styles(my_document)
    my_document.save('document.docx')
